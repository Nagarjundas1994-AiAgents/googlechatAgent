import os
import logging
from typing import List, Dict, Any
import fitz  # PyMuPDF
import docx
from bs4 import BeautifulSoup
from pdfminer.high_level import extract_text as pdfminer_extract_text

logger = logging.getLogger(__name__)

# Define chunk size (in tokens)
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200

async def process_document(file_path: str, filename: str) -> List[Dict[str, Any]]:
    """Process a document and extract text chunks
    
    Args:
        file_path: Path to the document file
        filename: Original filename
        
    Returns:
        List of text chunks with metadata
    """
    try:
        file_extension = os.path.splitext(filename)[1].lower()
        
        if file_extension == ".pdf":
            return await process_pdf(file_path, filename)
        elif file_extension == ".docx":
            return await process_docx(file_path, filename)
        elif file_extension == ".txt":
            return await process_txt(file_path, filename)
        elif file_extension in [".html", ".htm"]:
            return await process_html(file_path, filename)
        else:
            logger.warning(f"Unsupported file format: {file_extension}")
            return []
    except Exception as e:
        logger.error(f"Error processing document {filename}: {str(e)}")
        raise

async def process_pdf(file_path: str, filename: str) -> List[Dict[str, Any]]:
    """Process a PDF document and extract text chunks"""
    chunks = []
    
    try:
        # Method 1: Using PyMuPDF (faster but may have issues with some PDFs)
        doc = fitz.open(file_path)
        text = ""
        
        for page_num, page in enumerate(doc):
            text += page.get_text()
            
            # If we've accumulated enough text, create a chunk
            if len(text.split()) >= CHUNK_SIZE:
                chunks.append({
                    "text": text,
                    "metadata": {
                        "source": filename,
                        "page": page_num + 1
                    }
                })
                # Keep some overlap for context
                words = text.split()
                text = " ".join(words[-CHUNK_OVERLAP:]) if len(words) > CHUNK_OVERLAP else ""
        
        # Add any remaining text as a chunk
        if text.strip():
            chunks.append({
                "text": text,
                "metadata": {
                    "source": filename,
                    "page": doc.page_count
                }
            })
            
        doc.close()
        
        # If PyMuPDF failed to extract text, try pdfminer as a fallback
        if not chunks:
            logger.info(f"PyMuPDF extracted no text from {filename}, trying pdfminer as fallback")
            text = pdfminer_extract_text(file_path)
            
            # Split into chunks
            words = text.split()
            for i in range(0, len(words), CHUNK_SIZE - CHUNK_OVERLAP):
                chunk_words = words[i:i + CHUNK_SIZE]
                chunk_text = " ".join(chunk_words)
                
                chunks.append({
                    "text": chunk_text,
                    "metadata": {
                        "source": filename,
                        "chunk_index": i // (CHUNK_SIZE - CHUNK_OVERLAP)
                    }
                })
    except Exception as e:
        logger.error(f"Error processing PDF {filename}: {str(e)}")
        raise
    
    return chunks

async def process_docx(file_path: str, filename: str) -> List[Dict[str, Any]]:
    """Process a DOCX document and extract text chunks"""
    chunks = []
    
    try:
        doc = docx.Document(file_path)
        text = ""
        
        for para_index, para in enumerate(doc.paragraphs):
            text += para.text + "\n"
            
            # If we've accumulated enough text, create a chunk
            if len(text.split()) >= CHUNK_SIZE:
                chunks.append({
                    "text": text,
                    "metadata": {
                        "source": filename,
                        "paragraph_range": f"0-{para_index}"
                    }
                })
                # Keep some overlap for context
                words = text.split()
                text = " ".join(words[-CHUNK_OVERLAP:]) if len(words) > CHUNK_OVERLAP else ""
        
        # Add any remaining text as a chunk
        if text.strip():
            chunks.append({
                "text": text,
                "metadata": {
                    "source": filename,
                    "paragraph_range": f"0-{len(doc.paragraphs)}"
                }
            })
    except Exception as e:
        logger.error(f"Error processing DOCX {filename}: {str(e)}")
        raise
    
    return chunks

async def process_txt(file_path: str, filename: str) -> List[Dict[str, Any]]:
    """Process a TXT document and extract text chunks"""
    chunks = []
    
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        
        # Split into chunks
        words = text.split()
        for i in range(0, len(words), CHUNK_SIZE - CHUNK_OVERLAP):
            chunk_words = words[i:i + CHUNK_SIZE]
            chunk_text = " ".join(chunk_words)
            
            chunks.append({
                "text": chunk_text,
                "metadata": {
                    "source": filename,
                    "chunk_index": i // (CHUNK_SIZE - CHUNK_OVERLAP)
                }
            })
    except Exception as e:
        logger.error(f"Error processing TXT {filename}: {str(e)}")
        raise
    
    return chunks

async def process_html(file_path: str, filename: str) -> List[Dict[str, Any]]:
    """Process an HTML document and extract text chunks"""
    chunks = []
    
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            html_content = f.read()
        
        # Parse HTML
        soup = BeautifulSoup(html_content, "html.parser")
        
        # Remove script and style elements
        for script_or_style in soup(["script", "style"]):
            script_or_style.extract()
        
        # Get text
        text = soup.get_text()
        
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks_of_lines = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = "\n".join(chunk for chunk in chunks_of_lines if chunk)
        
        # Split into chunks
        words = text.split()
        for i in range(0, len(words), CHUNK_SIZE - CHUNK_OVERLAP):
            chunk_words = words[i:i + CHUNK_SIZE]
            chunk_text = " ".join(chunk_words)
            
            chunks.append({
                "text": chunk_text,
                "metadata": {
                    "source": filename,
                    "chunk_index": i // (CHUNK_SIZE - CHUNK_OVERLAP)
                }
            })
    except Exception as e:
        logger.error(f"Error processing HTML {filename}: {str(e)}")
        raise
    
    return chunks